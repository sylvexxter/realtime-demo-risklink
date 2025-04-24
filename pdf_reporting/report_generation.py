
import json
from docxtpl import DocxTemplate
from docx import Document
from docx2pdf import convert
from datetime import datetime
import os

# Paths
TEMPLATE_FILE = "risk_assessment_template.docx"
TEMP_RENDERED_FILE = "temp_rendered.docx"
FINAL_DOCX = "Risk_Assessment_Report.docx"
FINAL_PDF = "Risk_Assessment_Report.pdf"
USER_INFO_FILE = "../src/json_reports/initialRiskAssessment/user_information.json"
DOMAIN_FILES = [
    "../src/json_reports/initialRiskAssessment/training_awareness_report.json",
    "../src/json_reports/initialRiskAssessment/asset_management_report.json",
    "../src/json_reports/initialRiskAssessment/data_protection_privacy_report.json",
    "../src/json_reports/initialRiskAssessment/backups_report.json",
    "../src/json_reports/initialRiskAssessment/system_security_report.json",
    "../src/json_reports/initialRiskAssessment/anti_virus_report.json",
    "../src/json_reports/initialRiskAssessment/access_control_report.json",
    "../src/json_reports/initialRiskAssessment/incident_response_report.json"
]

LOGO_PATH = "logo.png"

# Load user info
with open(USER_INFO_FILE, "r") as f:
    raw_user = json.load(f)

user_info = {
    "name": raw_user.get("user_name", ""),
    "role": raw_user.get("user_position", ""),
    "organization": raw_user.get("user_company", "")
}

# Parse collection time
collection_ts = raw_user.get("collection_timestamp", "")
try:
    ts_obj = datetime.fromisoformat(collection_ts.replace("Z", "+00:00"))
except Exception as e:
    print(f"⚠️ Failed to parse timestamp: {e}")
    ts_obj = datetime.now()

context = {
    "date": ts_obj.strftime("%Y-%m-%d"),
    "time": ts_obj.strftime("%H:%M"),
    "assessor": user_info
}

# Render metadata fields
tpl = DocxTemplate(TEMPLATE_FILE)
tpl.render(context)
tpl.save(TEMP_RENDERED_FILE)

# Load rendered doc
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(TEMP_RENDERED_FILE)
if doc.tables:
    first_table = doc.tables[0]
    first_table._element.getparent().remove(first_table._element)
# Insert logo at the top (optional)
if os.path.exists(LOGO_PATH):
    doc.paragraphs[0].insert_paragraph_before().add_run().add_picture(LOGO_PATH, width=Inches(1.5))

# Add sectioned domain tables
for file in DOMAIN_FILES:
    file_title = os.path.basename(file).replace("_report.json", "").replace("_", " ").title()
    try:
        with open(file, "r") as f:
            data = json.load(f)
            if not data:
                continue

            doc.add_heading(file_title, level=2)

            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Question"
            hdr_cells[1].text = "Answer"
            hdr_cells[2].text = "Answer Context"

            for item in data:
                row = table.add_row().cells
                row[0].text = item.get("question_description", "")
                row[1].text = item.get("answer_text", "")
                row[2].text = item.get("answer_context", "")

    except FileNotFoundError:
        print(f"⚠️ Skipping missing file: {file}")
    except json.JSONDecodeError:
        print(f"⚠️ Skipping invalid JSON: {file}")

# Save output
company_clean = user_info["organization"].replace(" ", "_")
final_docx = f"Risk_Assessment_Report_{company_clean}.docx"
final_pdf = f"Risk_Assessment_Report_{company_clean}.pdf"
doc.save(final_docx)
convert(final_docx, final_pdf)

print(f"✅ PDF generated: {final_pdf}")
